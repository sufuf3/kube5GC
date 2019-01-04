#
# Copyright (c) 2017, NextEPC Group
# All rights reserved.
#
# Redistribution and use in source and binary forms, with or without
# modification, are permitted provided that the following conditions are met:
#
# 1. Redistributions of source code must retain the above copyright notice, this
#    list of conditions and the following disclaimer.
# 2. Redistributions in binary form must reproduce the above copyright notice,
#    this list of conditions and the following disclaimer in the documentation
#    and/or other materials provided with the distribution.
#
# THIS SOFTWARE IS PROVIDED BY THE COPYRIGHT HOLDERS AND CONTRIBUTORS "AS IS" AND
# ANY EXPRESS OR IMPLIED WARRANTIES, INCLUDING, BUT NOT LIMITED TO, THE IMPLIED
# WARRANTIES OF MERCHANTABILITY AND FITNESS FOR A PARTICULAR PURPOSE ARE
# DISCLAIMED. IN NO EVENT SHALL THE COPYRIGHT OWNER OR CONTRIBUTORS BE LIABLE FOR
# ANY DIRECT, INDIRECT, INCIDENTAL, SPECIAL, EXEMPLARY, OR CONSEQUENTIAL DAMAGES
# (INCLUDING, BUT NOT LIMITED TO, PROCUREMENT OF SUBSTITUTE GOODS OR SERVICES;
# LOSS OF USE, DATA, OR PROFITS; OR BUSINESS INTERRUPTION) HOWEVER CAUSED AND
# ON ANY THEORY OF LIABILITY, WHETHER IN CONTRACT, STRICT LIABILITY, OR TORT
# (INCLUDING NEGLIGENCE OR OTHERWISE) ARISING IN ANY WAY OUT OF THE USE OF THIS
# SOFTWARE, EVEN IF ADVISED OF THE POSSIBILITY OF SUCH DAMAGE.
#
from docx import Document
import re, os, sys, string
import datetime
import getopt
import getpass

version = "0.1.0"

msg_list = {}
type_list = {}
group_list = {}

verbosity = 0
filename = ""
outdir = './'
cachedir = './cache/'

FAIL = '\033[91m'
INFO = '\033[93m'
ENDC = '\033[0m'

def d_print(string):
    if verbosity > 0:
        sys.stdout.write(string)

def d_info(string):
    sys.stdout.write(INFO + string + ENDC + "\n")

def d_error(string):
    sys.stderr.write(FAIL + string + ENDC + "\n")
    sys.exit(0)

def write_file(f, string):
    f.write(string)
    d_print(string)

def output_header_to_file(f):
    now = datetime.datetime.now()
    f.write("""/*
 * Copyright (c) 2017, NextEPC Group
 * All rights reserved.
 *
 * Redistribution and use in source and binary forms, with or without
 * modification, are permitted provided that the following conditions are met:
 * 
 * 1. Redistributions of source code must retain the above copyright notice, this
 *    list of conditions and the following disclaimer.
 * 2. Redistributions in binary form must reproduce the above copyright notice,
 *    this list of conditions and the following disclaimer in the documentation
 *    and/or other materials provided with the distribution.

 * THIS SOFTWARE IS PROVIDED BY THE COPYRIGHT HOLDERS AND CONTRIBUTORS "AS IS" AND
 * ANY EXPRESS OR IMPLIED WARRANTIES, INCLUDING, BUT NOT LIMITED TO, THE IMPLIED
 * WARRANTIES OF MERCHANTABILITY AND FITNESS FOR A PARTICULAR PURPOSE ARE
 * DISCLAIMED. IN NO EVENT SHALL THE COPYRIGHT OWNER OR CONTRIBUTORS BE LIABLE FOR
 * ANY DIRECT, INDIRECT, INCIDENTAL, SPECIAL, EXEMPLARY, OR CONSEQUENTIAL DAMAGES
 * (INCLUDING, BUT NOT LIMITED TO, PROCUREMENT OF SUBSTITUTE GOODS OR SERVICES;
 * LOSS OF USE, DATA, OR PROFITS; OR BUSINESS INTERRUPTION) HOWEVER CAUSED AND
 * ON ANY THEORY OF LIABILITY, WHETHER IN CONTRACT, STRICT LIABILITY, OR TORT
 * (INCLUDING NEGLIGENCE OR OTHERWISE) ARISING IN ANY WAY OUT OF THE USE OF THIS
 * SOFTWARE, EVEN IF ADVISED OF THE POSSIBILITY OF SUCH DAMAGE.
 */

""")
    f.write("/*******************************************************************************\n")
    f.write(" * This file had been created by gtp_tlv.py script v%s\n" % (version))
    f.write(" * Please do not modify this file but regenerate it via script.\n")
    f.write(" * Created on: %s by %s\n * from %s\n" % (str(now), getpass.getuser(), filename))
    f.write(" ******************************************************************************/\n\n")

def usage():
    print "Python generating TLV build/parser for GTPv2-C v%s" % (version)
    print "Usage: python gtp_tlv.py [options]"
    print "Available options:"
    print "-d        Enable script debug"
    print "-f [file] Input file to parse"
    print "-o [dir]  Output files to given directory"
    print "-c [dir]  Cache files to given directory"
    print "-h        Print this help and return"

def v_upper(v):
    return re.sub('3GPP', '', re.sub('\'', '_', re.sub('/', '_', re.sub('-', '_', re.sub(' ', '_', v)))).upper())

def v_lower(v):
    return re.sub('3gpp', '', re.sub('\'', '_', re.sub('/', '_', re.sub('-', '_', re.sub(' ', '_', v)))).lower())

def get_cells(cells):
    #instance = cells[4].text.encode('ascii', 'ignore')
    #if instance.isdigit() is not True:
    #    return None
    instance = "0"  # PFCP has no instance
    note = cells[0].text.encode('ascii', 'ignore')
    if note.find('NOTE') != -1:
        return None
    comment = cells[2].text.encode('ascii', 'ignore')
    comment = re.sub('\n|\"|\'|\\\\', '', comment);
    #print comment
    ie_type = re.sub('\s*$', '', re.sub('\'\s*\n*\s*\(NOTE.*\)*', '', cells[-1].text.encode('ascii', 'ignore')))    
    
    #if ie_type.find('Usage Report') != -1:
    if ie_type == 'Usage Report':
        if comment.find('Report Type') != -1:
            ie_type = "Usage Report Session Report Request"
        elif comment.find('Query URR') != -1:
            ie_type = "Usage Report in Session Modification Response"
        elif comment.find('provisioned ') != -1:
            ie_type = "Usage Report Session Deletion Response"
        else:
             assert False, "Unknown IE type : [Usage Report]"
    
    if ie_type == 'Update BAR':
        if comment.find('7.5.4.3-3') != -1:
            ie_type = "Update BAR Session Modification Request"
        elif comment.find('7.5.9.2-1') != -1:
            ie_type = "Update BAR PFCP Session Report Response"
        else:
             assert False, "Unknown IE type : [Update BAR]"
    
    if ie_type == 'Metric' or ie_type == 'FQ-CSID' or ie_type == 'PDN Type':
        ie_type = ie_type + 'p'        
    elif ie_type == 'Load Control Information' or ie_type == 'Overload Control Information': 
        ie_type = ie_type + 'p'
        
    if ie_type.find('PFD Contents') != -1:
        ie_type = 'PFD contents'
    elif ie_type.find('PFD') != -1:
        ie_type = 'PFD context'
    elif ie_type.find('PDR ID') != -1:
        ie_type = 'Packet Detection Rule ID'
    elif ie_type.find('UE IP address') != -1:
        ie_type = 'UE IP Address'
    elif ie_type.find('SxSMReq-Flags') != -1:
        ie_type = 'PFCPSMReq-Flags'
    elif ie_type.find('PFCPSRRsp-Flags2') != -1:
        ie_type = 'PFCPSRRsp-Flags'
    elif ie_type.find('IPv4 Configuration Parameters (IP4CP)') != -1:
        ie_type = 'IP4CP'
    if ie_type not in type_list.keys():
        assert False, "Unknown IE type : [" \
                + cells[-1].text + "]" + "(" + ie_type + ")"
    presence = cells[1].text.encode('ascii', 'ignore')
    ie_value = re.sub('\s*\n*\s*\([^\)]*\)*', '', cells[0].text).encode('ascii', 'ignore')
    if ie_value[len(ie_value)-1] == ' ':
        ie_value = ie_value[:len(ie_value)-1]

    # 0402: ?
    #print ie_type
    if ie_type == 'Create PDR' or ie_type == 'Create FAR' or ie_type == 'Update PDR':
        instance = "1"
        print ie_type + ' ' + type_list[ie_type]["max_instance"] + '\n'

    if int(instance) > int(type_list[ie_type]["max_instance"]):
        type_list[ie_type]["max_instance"] = instance
        write_file(f, "type_list[\"" + ie_type + "\"][\"max_instance\"] = \"" + instance + "\"\n")
        print "CCCCCCCCCCCCCCCCCCCCCCCCCCCC\n"

    return { "ie_type" : ie_type, "ie_value" : ie_value, "presence" : presence, "instance" : instance, "comment" : comment }

def write_cells_to_file(name, cells):
    write_file(f, name + ".append({ \"ie_type\" : \"" + cells["ie_type"] + \
        "\", \"ie_value\" : \"" + cells["ie_value"] + \
        "\", \"presence\" : \"" + cells["presence"] + \
        "\", \"instance\" : \"" + cells["instance"] + \
        "\", \"comment\" : \"" + cells["comment"] + "\"})\n")

try:
    opts, args = getopt.getopt(sys.argv[1:], "df:ho:c:", ["debug", "file", "help", "output", "cache"])
except getopt.GetoptError as err:
    # print help information and exit:
    usage()
    sys.exit(2)

for o, a in opts:
    if o in ("-d", "--debug"):
        verbosity = 1
    if o in ("-f", "--file"):
        filename = a
    if o in ("-o", "--output"):
        outdir = a
        if outdir.rfind('/') != len(outdir):
            outdir += '/'
    if o in ("-c", "--cache"):
        cache = a
        if cachedir.rfind('/') != len(cachedir):
            cachedir += '/'
    if o in ("-h", "--help"):
        usage()
        sys.exit(2)

if os.path.isfile(filename) and os.access(filename, os.R_OK):
    file = open(filename, 'r') 
else:
    d_error("Cannot find file : " + filename)

d_info("[Message List]")
cachefile = cachedir + 'tlv_msg_list.py'
if os.path.isfile(cachefile) and os.access(cachefile, os.R_OK):
    execfile(cachefile)
    print "Read from " + cachefile
else:
    document = Document(filename)
    f = open(cachefile, 'w') 

    msg_table = ""
    for i, table in enumerate(document.tables):
        cell = table.rows[0].cells[0]
        if cell.text.find('Message Type value') != -1:
            msg_table = table
            d_print("Table Index = %d\n" % i)
            #print("Table Index = %d\n" % i)

    for row in msg_table.rows[2:-4]:
        key = row.cells[1].text.encode('ascii', 'ignore')
        type = row.cells[0].text.encode('ascii', 'ignore')
        if type.isdigit() is False:
            continue
        if int(type) in range(128, 160):
            continue
        if int(type) in range(231, 240):
            continue
        if key.find('Reserved') != -1:
            continue
        key = re.sub('\s*\n*\s*\([^\)]*\)*', '', key)
        msg_list[key] = { "type": type }
        #print("msg type =" + type)
        write_file(f, "msg_list[\"" + key + "\"] = { \"type\" : \"" + type + "\" }\n")
    f.close()

d_info("[IE Type List]")
cachefile = cachedir + 'tlv_type_list.py'
if os.path.isfile(cachefile) and os.access(cachefile, os.R_OK):
    execfile(cachefile)
    print "Read from " + cachefile
else:
    document = Document(filename)
    f = open(cachefile, 'w') 

    ie_table = ""
    for i, table in enumerate(document.tables):
        cell = table.rows[0].cells[0]
        if cell.text.find('IE Type value') != -1:
            ie_table = table
            d_print("Table Index = %d\n" % i)

    for row in ie_table.rows[1:-1]:
        key = row.cells[1].text.encode('ascii', 'ignore')
        if key.find('Reserved') != -1:
            continue
        if key.find('MM Context') != -1:
            continue
        #elif key.find('Recovery') != -1:
        #    key = 'Recovery'
        elif key.find('Trusted WLAN Mode Indication') != -1:
            key = 'TWMI'
        elif key.find('LDN') != -1:
            key = 'LDN'
        elif key.find('APCO') != -1:
            key = 'APCO'
        elif key.find('Remote UE IP information') != -1:
            key = 'Remote UE IP Information'
        elif key.find('Procedure Transaction ID') != -1:
            key = 'PTI'
        else:
            #key = re.sub('.*\(', '', row.cells[1].text.encode('ascii', 'ignore'))
            key = re.sub('\(', '', key)
            key = re.sub('\)', '', key)
            key = re.sub('\s*$', '', key)
        type = row.cells[0].text.encode('ascii', 'ignore')
        if key == 'Metric' or key == 'FQ-CSID' or key == 'PDN Type':
            key = key + 'p'
        elif key == 'Load Control Information' or key == 'Overload Control Information': 
            key = key + 'p'    
        # 0402: here is active !
        #if key == 'Create PDR' or key == 'Create FAR' or key == 'Update PDR':
        #    type_list[key] = { "type": type , "max_instance" : "1" }
        #else:
        type_list[key] = { "type": type , "max_instance" : "0" }
        write_file(f, "type_list[\"" + key + "\"] = { \"type\" : \"" + type)
        write_file(f, "\", \"max_instance\" : \"0\" }\n")
    f.close()
#type_list['MM Context'] = { "type": "107", "max_instance" : "0" }

d_info("[Group IE List]")
cachefile = cachedir + 'tlv_group_list.py'
if os.path.isfile(cachefile) and os.access(cachefile, os.R_OK):
    execfile(cachefile)
    print "Read from " + cachefile
else:
    document = Document(filename)
    f = open(cachefile, 'w') 

    for i, table in enumerate(document.tables):
        if table.rows[0].cells[0].text.find('Octet') != -1 and \
            table.rows[0].cells[1].text.find('Outer Header to be created') == -1 and \
            table.rows[0].cells[2].text.find('IE Type') != -1:
            d_print("Table Index = %d\n" % i)

            row = table.rows[0];

            if len(re.findall('\d+', row.cells[2].text)) == 0:
                continue;
            ie_type = re.findall('\d+', row.cells[2].text)[0].encode('ascii', 'ignore')
            ie_name = re.sub('\s*IE Type.*', '', row.cells[2].text.encode('ascii', 'ignore'))

            if (int(ie_type) == 78):
                ie_name =  "Usage Report in Session Modification Response"
            elif (int(ie_type) == 79):
                ie_name =  "Usage Report Session Deletion Response"
            elif (int(ie_type) == 80):
                ie_name =  "Usage Report Session Report Request"    
            elif (int(ie_type) == 86):
                ie_name =  "Update BAR Session Modification Request" 
            elif (int(ie_type) == 12):
                ie_name =  "Update BAR PFCP Session Report Response" 
            
            if ie_name == 'PFD':
                ie_name = 'PFD context'
            
            if ie_name == 'Load Control Information' or ie_name == 'Overload Control Information': 
                ie_name = ie_name + 'p' 

            if ie_name not in group_list.keys():
                ies = []
                write_file(f, "ies = []\n")
                for row in table.rows[4:]:
                    cells = get_cells(row.cells)
                    if cells is None:
                        continue
                    #if cells["ie_type"] == 'Create PDR' or cells["ie_type"] == 'Create FAR' or cells["ie_type"] == 'Update PDR':
                    #    cells["instance"] = str(int(cells["instance"])+1)
                    #    print 'Hahahaha-----------------------\n'

                    ies_is_added = True
                    for ie in ies:
                        if (cells["ie_type"], cells["instance"]) == (ie["ie_type"], ie["instance"]):
                            ies_is_added = False
                    if ies_is_added is True:
                        ies.append(cells)
                        write_cells_to_file("ies", cells)

                group_list[ie_name] = { "type" : ie_type, "ies" : ies }
                write_file(f, "group_list[\"" + ie_name + "\"] = { \"type\" : \"" + ie_type + "\", \"ies\" : ies }\n")
            else:
                group_list_is_added = False
                added_ies = group_list[ie_name]["ies"]
                write_file(f, "added_ies = group_list[\"" + ie_name + "\"][\"ies\"]\n")
                for row in table.rows[4:]:
                    cells = get_cells(row.cells)
                    if cells is None:
                        continue

                    ies_is_added = True
                    for ie in group_list[ie_name]["ies"]:
                        if (cells["ie_type"], cells["instance"]) == (ie["ie_type"], ie["instance"]):
                            ies_is_added = False
                    for ie in ies:
                        if (cells["ie_type"], cells["instance"]) == (ie["ie_type"], ie["instance"]):
                            ies_is_added = False
                    if ies_is_added is True:
                        added_ies.append(cells)
                        write_cells_to_file("added_ies", cells)
                        group_list_is_added = True
                if group_list_is_added is True:
                    group_list[ie_name] = { "type" : ie_type, "ies" : added_ies }
                    write_file(f, "group_list[\"" + ie_name + "\"] = { \"type\" : \"" + ie_type + "\", \"ies\" : added_ies }\n")
    f.close()

msg_list["PFCP Heartbeat Request"]["table"] = 6
msg_list["PFCP Heartbeat Response"]["table"] = 7
msg_list["PFCP Association Setup Request"]["table"] = 12
msg_list["PFCP Association Setup Response"]["table"] = 13
msg_list["PFCP Association Update Request"]["table"] = 14
msg_list["PFCP Association Update Response"]["table"] = 15
msg_list["PFCP Association Release Request"]["table"] = 16
msg_list["PFCP Association Release Response"]["table"] = 17
msg_list["PFCP Version Not Supported Response"]["table"] = 11
msg_list["PFCP Node Report Request"]["table"] = 18
msg_list["PFCP Node Report Response"]["table"] = 20
msg_list["PFCP Session Set Deletion Request"]["table"] = 21
msg_list["PFCP Session Set Deletion Response"]["table"] = 22
msg_list["PFCP Session Establishment Request"]["table"] = 23
msg_list["PFCP Session Establishment Response"]["table"] = 33
msg_list["PFCP Session Modification Request"]["table"] = 37
msg_list["PFCP Session Modification Response"]["table"] = 51
msg_list["PFCP Session Deletion Request"]["table"] = 53
msg_list["PFCP Session Deletion Response"]["table"] = 54
msg_list["PFCP Session Report Request"]["table"] = 56
msg_list["PFCP Session Report Response"]["table"] = 57

for key in msg_list.keys():
    if "table" in msg_list[key].keys():
        d_info("[" + key + "]")
        cachefile = cachedir + "tlv_msg_" + msg_list[key]["type"] + ".py"
        if os.path.isfile(cachefile) and os.access(cachefile, os.R_OK):
            execfile(cachefile)
            print "Read from " + cachefile
        else:
            document = Document(filename)
            f = open(cachefile, 'w') 

            ies = []
            write_file(f, "ies = []\n")
            table = document.tables[msg_list[key]["table"]]
            if key.find('Association') != -1:
                start_i = 1
            elif key.find('Heartbeat') != -1:
                start_i = 1
            else:
                start_i = 2
            
            if key != "PFCP Session Deletion Request":   # this msg is null     
                for row in table.rows[start_i:]:
                    cells = get_cells(row.cells)
                    if cells is None:
                        continue
    
                    # 0402
                    if (cells["ie_type"] == 'Create PDR' or cells["ie_type"] == 'Create FAR' or cells["ie_type"] == 'Update PDR'):
                        cells["instance"] = '0' 
                        cells["presence"] = 'O'
                        ies.append(cells)
                        write_cells_to_file("ies", cells)
                    cells = get_cells(row.cells)
    
                    ies_is_added = True
                    for ie in ies:
                        #0403 modify
                        if (cells["ie_type"], cells["instance"]) == (ie["ie_type"], ie["instance"]):
                        #if (cells["ie_type"], cells["ie_value"]) == (ie["ie_type"], ie["ie_value"]):
                            ies_is_added = False
                    if ies_is_added is True:
                        ies.append(cells)
                        write_cells_to_file("ies", cells)
                    # 0402
                    #if (cells["ie_type"] == 'Create PDR' or cells["ie_type"] == 'Create FAR' or cells["ie_type"] == 'Update PDR'):
                    #    cells["instance"] = str(int(cells["instance"])+1)
                    #    cells["ie_value"] = cells["ie_value"] + '1' 
                    #    cells["presence"] ='O'
                    #    ies.append(cells)
                    #    write_cells_to_file("ies", cells)
            msg_list[key]["ies"] = ies
            write_file(f, "msg_list[key][\"ies\"] = ies\n")
            f.close()

#type_list["Recovery"]["size"] = 1                       # Type : 3
#type_list["EBI"]["size"] = 1                            # Type : 73
#type_list["RAT Type"]["size"] = 1                       # Type : 82
#type_list["PDN Type"]["size"] = 1                       # Type : 99
#type_list["Port Number"]["size"] = 2                    # Type : 126
#type_list["APN Restriction"]["size"] = 1                # Type : 127
#type_list["Selection Mode"]["size"] = 1                 # Type : 128
#type_list["Node Type"]["size"] = 1                 # Type : 128

f = open(outdir + 'pfcp_message.h', 'w')
output_header_to_file(f)
f.write("""#ifndef __PFCP_TLV_H__
#define __PFCP_TLV_H__

#include "core_tlv_msg.h"

#ifdef __cplusplus
extern "C" {
#endif /* __cplusplus */

/* 5.1 General format */
#define PFCP_HEADER_LEN     16
#define PFCP_SEID_LEN        8
typedef struct _pfcp_header_t {
    union {
        struct {
        ED4(c_uint8_t version:3;,
            c_uint8_t spare1:3;,
            c_uint8_t mp:1;,
            c_uint8_t seid_p:1;)
        };
        c_uint8_t flags;
    };
    c_uint8_t type;
    c_uint16_t length;
    union {
        struct {
            c_uint64_t seid;
            /* sqn : 31bit ~ 8bit, spare : 7bit ~ 0bit */
#define PFCP_XID_TO_SQN(__xid) htonl(((__xid) << 8))
#define PFCP_SQN_TO_XID(__sqn) (ntohl(__sqn) >> 8)
            c_uint32_t sqn;
        };
        /* sqn : 31bit ~ 8bit, spare : 7bit ~ 0bit */
        c_uint32_t sqn_only;
    };
} __attribute__ ((packed)) pfcp_header_t;



/* PFCP message type */
""")

tmp = [(k, v["type"]) for k, v in msg_list.items()]
sorted_msg_list = sorted(tmp, key=lambda tup: int(tup[1]))
for (k, v) in sorted_msg_list:
    f.write("#define " + v_upper(k) + "_TYPE " + v + "\n")
f.write("\n")

tmp = [(k, v["type"]) for k, v in type_list.items()]
sorted_type_list = sorted(tmp, key=lambda tup: int(tup[1]))
for (k, v) in sorted_type_list:
    f.write("#define TLV_PFCP_" + v_upper(k) + "_TYPE " + v + "\n")
f.write("\n")

f.write("/* Infomration Element TLV Descriptor */\n")
for (k, v) in sorted_type_list:
    if k in group_list.keys():
        continue
    for instance in range(0, int(type_list[k]["max_instance"])+1):
        if v_lower(k)=="cause" or v_lower(k)=="sequence_number" or v_lower(k)=="f_teid":
            f.write("extern tlv_desc_t tlv_desc_pfcp" + v_lower(k))
        else:
            f.write("extern tlv_desc_t tlv_desc_" + v_lower(k))
        f.write("_" + str(instance) + ";\n")
f.write("\n")

tmp = [(k, v["type"]) for k, v in group_list.items()]
sorted_group_list = sorted(tmp, key=lambda tup: int(tup[1]), reverse=False)

f.write("/* Group Infomration Element TLV Descriptor */\n")
for (k, v) in sorted_group_list:
    for instance in range(0, int(type_list[k]["max_instance"])+1):
        f.write("extern tlv_desc_t tlv_desc_" + v_lower(k))
        f.write("_" + str(instance) + ";\n")
f.write("\n")

f.write("/* Message Descriptor */\n")
for (k, v) in sorted_msg_list:
    f.write("extern tlv_desc_t tlv_desc_" + v_lower(k) + ";\n")
f.write("\n")

f.write("/* Structure for Infomration Element */\n")
for (k, v) in sorted_type_list:
    if k in group_list.keys():
        continue
    if "size" in type_list[k]:
        if type_list[k]["size"] == 1:
            f.write("typedef tlv_uint8_t tlv_" + v_lower(k) + "_t;\n")
        elif type_list[k]["size"] == 2:
            f.write("typedef tlv_uint16_t tlv_" + v_lower(k) + "_t;\n")
        elif type_list[k]["size"] == 3:
            f.write("typedef tlv_uint24_t tlv_" + v_lower(k) + "_t;\n")
        elif type_list[k]["size"] == 4:
            f.write("typedef tlv_uint32_t tlv_" + v_lower(k) + "_t;\n")
        else:
            assert False, "Unknown size = %d for key = %s" % (type_list[k]["size"], k)
    else:
        f.write("typedef tlv_octet_t tlv_" + v_lower(k) + "_t;\n")
f.write("\n")

#for (k, v) in sorted_group_list:
#    f.write("typedef struct _tlv_" + v_lower(k) + "_t ")
#    f.write(" tlv_" + v_lower(k) + "_t;\n")
#    f.write("\n")
tmp = []
f.write("/* Structure for Group Infomration Element */\n")
for (k, v) in sorted_group_list:
    if v_lower(k) == "create_pdr":
        tmp.append(k)
        continue 
    if v_lower(k) == "create_far":
        tmp.append(k)
        continue
    if v_lower(k) == "update_far":
        tmp.append(k)
        continue
    if v_lower(k) == "application_id_s_pfds":
        tmp.append(k)
        continue    
    f.write("typedef struct _tlv_" + v_lower(k) + "_t {\n")
    f.write("    tlv_presence_t presence;\n")
    for ies in group_list[k]["ies"]:
        f.write("    tlv_" + v_lower(ies["ie_type"]) + "_t " + \
                v_lower(ies["ie_value"]))
        if ies["ie_type"] == "F-TEID":
            if ies["ie_value"] == "S2b-U ePDG F-TEID":
                f.write("_" + ies["instance"] + ";")
            elif ies["ie_value"] == "S2a-U TWAN F-TEID":
                f.write("_" + ies["instance"] + ";")
            else:
                f.write(";")
            f.write(" /* Instance : " + ies["instance"] + " */\n")
        else:
            f.write(";\n")
    f.write("} tlv_" + v_lower(k) + "_t; ;\n")
    f.write("\n")

for k in tmp:
    f.write("typedef struct _tlv_" + v_lower(k) + "_t {\n")
    f.write("    tlv_presence_t presence;\n")
    for ies in group_list[k]["ies"]:
        f.write("    tlv_" + v_lower(ies["ie_type"]) + "_t " + \
                v_lower(ies["ie_value"]))
        if ies["ie_type"] == "F-TEID":
            if ies["ie_value"] == "S2b-U ePDG F-TEID":
                f.write("_" + ies["instance"] + ";")
            elif ies["ie_value"] == "S2a-U TWAN F-TEID":
                f.write("_" + ies["instance"] + ";")
            else:
                f.write(";")
            f.write(" /* Instance : " + ies["instance"] + " */\n")
        else:
            f.write(";\n")
    f.write("} tlv_" + v_lower(k) + "_t; ;\n")
    f.write("\n")

f.write("/* Structure for Message */\n")
for (k, v) in sorted_msg_list:
    if "ies" in msg_list[k]:
        f.write("typedef struct _" + v_lower(k) + "_t {\n")
        for ies in msg_list[k]["ies"]:
            # 0403 modify
            if ies["instance"] != "0":
                f.write("    tlv_" + v_lower(ies["ie_type"]) + "_t " + \
                    v_lower(ies["ie_value"]) + ies["instance"] + ";\n")
            else:
                f.write("    tlv_" + v_lower(ies["ie_type"]) + "_t " + \
                    v_lower(ies["ie_value"]) + ";\n")
        f.write("} " + v_lower(k) + "_t;\n")
        f.write("\n")

f.write("typedef struct _pfcp_message_t {\n")
f.write("   pfcp_header_t h;\n")
f.write("   union {\n")
for (k, v) in sorted_msg_list:
    if "ies" in msg_list[k]:
        f.write("        " + v_lower(k) + "_t " + v_lower(k) + ";\n");
f.write("   };\n");
f.write("} pfcp_message_t;\n\n")

f.write("""CORE_DECLARE(status_t) pfcp_parse_msg(
        pfcp_message_t *pfcp_message, pkbuf_t *pkbuf);
CORE_DECLARE(status_t) pfcp_build_msg(
        pkbuf_t **pkbuf, pfcp_message_t *pfcp_message);

#ifdef __cplusplus
}
#endif /* __cplusplus */

#endif /* __PFCP_TLV_H__ */
""")
f.close()

f = open(outdir + 'pfcp_message.c', 'w')
output_header_to_file(f)
f.write("""#define TRACE_MODULE _pfcp_message
#include "core_debug.h"
#include "pfcp_message.h"

""")

for (k, v) in sorted_type_list:
    if k in group_list.keys():
        continue
    for instance in range(0, int(type_list[k]["max_instance"])+1):
        if v_lower(k)=="cause" or v_lower(k)=="sequence_number" or v_lower(k)=="f_teid":
            f.write("tlv_desc_t tlv_desc_pfcp%s_%d =\n" % (v_lower(k), instance))
        else:
            f.write("tlv_desc_t tlv_desc_%s_%d =\n" % (v_lower(k), instance))
        f.write("{\n")
        if "size" in type_list[k]:
            if type_list[k]["size"] == 1:
                f.write("    TLV_UINT8,\n")
            elif type_list[k]["size"] == 2:
                f.write("    TLV_UINT16,\n")
            elif type_list[k]["size"] == 3:
                f.write("    TLV_UINT24,\n")
            elif type_list[k]["size"] == 4:
                f.write("    TLV_UINT32,\n")
            else:
                assert False, "Unknown size = %d for key = %s" % (type_list[k]["size"], k)
        else:
            f.write("    TLV_VAR_STR,\n")
        f.write("    \"%s\",\n" % k)
        f.write("    TLV_PFCP_%s_TYPE,\n" % v_upper(k))
        if "size" in type_list[k]:
            f.write("    %d,\n" % type_list[k]["size"])
        else:
            f.write("    0,\n")
        f.write("    %d,\n" % instance)
        f.write("    sizeof(tlv_%s_t),\n" % v_lower(k))
        f.write("    { NULL }\n")
        f.write("};\n\n")

for (k, v) in sorted_group_list:
    for instance in range(0, int(type_list[k]["max_instance"])+1):
        f.write("tlv_desc_t tlv_desc_%s_%d =\n" % (v_lower(k), instance))
        f.write("{\n")
        f.write("    TLV_COMPOUND,\n")
        f.write("    \"%s\",\n" % k)
        f.write("    TLV_PFCP_%s_TYPE,\n" % v_upper(k))
        f.write("    0,\n")
        f.write("    %d,\n" % instance)
        f.write("    sizeof(tlv_%s_t),\n" % v_lower(k))
        f.write("    {\n")
        for ies in group_list[k]["ies"]:
                if v_lower(ies["ie_type"])=="cause" or v_lower(ies["ie_type"])=="sequence_number" or v_lower(ies["ie_type"])=="f_teid":
                    f.write("        &tlv_desc_pfcp%s_%s,\n" % (v_lower(ies["ie_type"]), v_lower(ies["instance"])))
                else:
                    f.write("        &tlv_desc_%s_%s,\n" % (v_lower(ies["ie_type"]), v_lower(ies["instance"])))
        f.write("        NULL,\n")
        f.write("    }\n")
        f.write("};\n\n")

for (k, v) in sorted_msg_list:
    if "ies" in msg_list[k]:
        f.write("tlv_desc_t tlv_desc_%s =\n" % v_lower(k))
        f.write("{\n")
        f.write("    TLV_MESSAGE,\n")
        f.write("    \"%s\",\n" % k)
        f.write("    0, 0, 0, 0, {\n")
        for ies in msg_list[k]["ies"]:
                if v_lower(ies["ie_type"])=="cause" or v_lower(ies["ie_type"])=="sequence_number" or v_lower(ies["ie_type"])=="f_teid":
                    f.write("        &tlv_desc_pfcp%s_%s,\n" % (v_lower(ies["ie_type"]), v_lower(ies["instance"])))
                else:
                    f.write("        &tlv_desc_%s_%s,\n" % (v_lower(ies["ie_type"]), v_lower(ies["instance"])))
        f.write("    NULL,\n")
        f.write("}};\n\n")
f.write("\n")

f.write("""status_t pfcp_parse_msg(pfcp_message_t *pfcp_message, pkbuf_t *pkbuf)
{
    status_t rv = CORE_ERROR;
    pfcp_header_t *h = NULL;
    c_uint16_t size = 0;

    d_assert(pfcp_message, return CORE_ERROR, "Null param");
    d_assert(pkbuf, return CORE_ERROR, "Null param");
    d_assert(pkbuf->payload, return CORE_ERROR, "Null param");

    d_trace(50, "[PFCP] RECV : ");
    d_trace_hex(50, pkbuf->payload, pkbuf->len);

    h = pkbuf->payload;
    d_assert(h, return CORE_ERROR, "Null param");
    
    memset(pfcp_message, 0, sizeof(pfcp_message_t));

    if (h->seid_p)
        size = PFCP_HEADER_LEN;
    else
        size = PFCP_HEADER_LEN-PFCP_SEID_LEN;

    d_assert(pkbuf_header(pkbuf, -size) == CORE_OK,
            return CORE_ERROR, "pkbuf_header error");
    memcpy(&pfcp_message->h, pkbuf->payload - size, size);

    if (h->seid_p)
        pfcp_message->h.seid = be64toh(pfcp_message->h.seid);
    else
    {
        pfcp_message->h.sqn  = pfcp_message->h.sqn_only;
        pfcp_message->h.sqn_only  = pfcp_message->h.sqn_only;
    }

    if (pkbuf->len == 0)
        return CORE_OK;

    switch(pfcp_message->h.type)
    {
""")
for (k, v) in sorted_msg_list:
    if "ies" in msg_list[k]:
        f.write("        case %s_TYPE:\n" % v_upper(k))
        f.write("            rv = tlv_parse_msg(&pfcp_message->%s,\n" % v_lower(k))
        f.write("                    &tlv_desc_%s, pkbuf, TLV_MODE_T2_L2);\n" % v_lower(k))
        f.write("            break;\n")
f.write("""        default:
            d_warn("Not implmeneted(type:%d)", pfcp_message->h.type);
            break;
    }

    return rv;
}

""")

f.write("""status_t pfcp_build_msg(pkbuf_t **pkbuf, pfcp_message_t *pfcp_message)
{
    status_t rv = CORE_ERROR;

    d_assert(pfcp_message, return rv, "Null param");
    switch(pfcp_message->h.type)
    {
""")
for (k, v) in sorted_msg_list:
    if "ies" in msg_list[k]:
        f.write("        case %s_TYPE:\n" % v_upper(k))
        f.write("            rv = tlv_build_msg(pkbuf, &tlv_desc_%s,\n" % v_lower(k))
        f.write("                    &pfcp_message->%s, TLV_MODE_T2_L2);\n" % v_lower(k))
        f.write("            break;\n")
f.write("""        default:
            d_warn("Not implmeneted(type:%d)", pfcp_message->h.type);
            break;
    }

    if ((*pkbuf) && (*pkbuf)->payload)
    {
        d_trace(50, "[PFCP] SEND : ");
        d_trace_hex(50, (*pkbuf)->payload, (*pkbuf)->len);
    }

    return rv;
}
""")

f.write("\n")

f.close()
