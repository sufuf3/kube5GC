#
# Copyright (c) 2017, NextEPC Group
# All rights reserved.
#
# Redistribution and use in source and binary forms, with or without
# modification, are permitted provided that the following conditions are met:
#
# 1. Redistributions of source code must retain the above copyright notice, this
#    list of conditions and the following disclaimer.
# 2. Redistributions in binary form must reproduce the above copyright notice,
#    this list of conditions and the following disclaimer in the documentation
#    and/or other materials provided with the distribution.
#
# THIS SOFTWARE IS PROVIDED BY THE COPYRIGHT HOLDERS AND CONTRIBUTORS "AS IS" AND
# ANY EXPRESS OR IMPLIED WARRANTIES, INCLUDING, BUT NOT LIMITED TO, THE IMPLIED
# WARRANTIES OF MERCHANTABILITY AND FITNESS FOR A PARTICULAR PURPOSE ARE
# DISCLAIMED. IN NO EVENT SHALL THE COPYRIGHT OWNER OR CONTRIBUTORS BE LIABLE FOR
# ANY DIRECT, INDIRECT, INCIDENTAL, SPECIAL, EXEMPLARY, OR CONSEQUENTIAL DAMAGES
# (INCLUDING, BUT NOT LIMITED TO, PROCUREMENT OF SUBSTITUTE GOODS OR SERVICES;
# LOSS OF USE, DATA, OR PROFITS; OR BUSINESS INTERRUPTION) HOWEVER CAUSED AND
# ON ANY THEORY OF LIABILITY, WHETHER IN CONTRACT, STRICT LIABILITY, OR TORT
# (INCLUDING NEGLIGENCE OR OTHERWISE) ARISING IN ANY WAY OUT OF THE USE OF THIS
# SOFTWARE, EVEN IF ADVISED OF THE POSSIBILITY OF SUCH DAMAGE.
#
from docx import Document
import re, os, sys, string
import datetime
import getopt
import getpass

version = "0.1.0"

msg_list = {}
type_list = {}
group_list = {}

verbosity = 0
filename = ""
outdir = './'
cachedir = './cache/'

FAIL = '\033[91m'
INFO = '\033[93m'
ENDC = '\033[0m'

def d_print(string):
    if verbosity > 0:
        sys.stdout.write(string)

def d_info(string):
    sys.stdout.write(INFO + string + ENDC + "\n")

def d_error(string):
    sys.stderr.write(FAIL + string + ENDC + "\n")
    sys.exit(0)

def write_file(f, string):
    f.write(string)
    d_print(string)

def output_header_to_file(f):
    now = datetime.datetime.now()
    f.write("""/*
 * Copyright (c) 2017, NextEPC Group
 * All rights reserved.
 *
 * Redistribution and use in source and binary forms, with or without
 * modification, are permitted provided that the following conditions are met:
 * 
 * 1. Redistributions of source code must retain the above copyright notice, this
 *    list of conditions and the following disclaimer.
 * 2. Redistributions in binary form must reproduce the above copyright notice,
 *    this list of conditions and the following disclaimer in the documentation
 *    and/or other materials provided with the distribution.

 * THIS SOFTWARE IS PROVIDED BY THE COPYRIGHT HOLDERS AND CONTRIBUTORS "AS IS" AND
 * ANY EXPRESS OR IMPLIED WARRANTIES, INCLUDING, BUT NOT LIMITED TO, THE IMPLIED
 * WARRANTIES OF MERCHANTABILITY AND FITNESS FOR A PARTICULAR PURPOSE ARE
 * DISCLAIMED. IN NO EVENT SHALL THE COPYRIGHT OWNER OR CONTRIBUTORS BE LIABLE FOR
 * ANY DIRECT, INDIRECT, INCIDENTAL, SPECIAL, EXEMPLARY, OR CONSEQUENTIAL DAMAGES
 * (INCLUDING, BUT NOT LIMITED TO, PROCUREMENT OF SUBSTITUTE GOODS OR SERVICES;
 * LOSS OF USE, DATA, OR PROFITS; OR BUSINESS INTERRUPTION) HOWEVER CAUSED AND
 * ON ANY THEORY OF LIABILITY, WHETHER IN CONTRACT, STRICT LIABILITY, OR TORT
 * (INCLUDING NEGLIGENCE OR OTHERWISE) ARISING IN ANY WAY OUT OF THE USE OF THIS
 * SOFTWARE, EVEN IF ADVISED OF THE POSSIBILITY OF SUCH DAMAGE.
 */

""")
    f.write("/*******************************************************************************\n")
    f.write(" * This file had been created by gtp_tlv.py script v%s\n" % (version))
    f.write(" * Please do not modify this file but regenerate it via script.\n")
    f.write(" * Created on: %s by %s\n * from %s\n" % (str(now), getpass.getuser(), filename))
    f.write(" ******************************************************************************/\n\n")

def usage():
    print "Python generating TLV build/parser for GTPv2-C v%s" % (version)
    print "Usage: python gtp_tlv.py [options]"
    print "Available options:"
    print "-d        Enable script debug"
    print "-f [file] Input file to parse"
    print "-o [dir]  Output files to given directory"
    print "-c [dir]  Cache files to given directory"
    print "-h        Print this help and return"

def v_upper(v):
    return re.sub('3GPP', '', re.sub('\'', '_', re.sub('/', '_', re.sub('-', '_', re.sub(' ', '_', v)))).upper())

def v_lower(v):
    return re.sub('3gpp', '', re.sub('\'', '_', re.sub('/', '_', re.sub('-', '_', re.sub(' ', '_', v)))).lower())

def get_cells(cells):
    instance = cells[4].text.encode('ascii', 'ignore')
    if instance.isdigit() is not True:
        return None
    ie_type = re.sub('\s*$', '', re.sub('\s*\n*\s*\(NOTE.*\)*', '', cells[3].text.encode('ascii', 'ignore')))
    if ie_type.find('LDN') != -1:
        ie_type = 'LDN'
    elif ie_type.find('APCO') != -1:
        ie_type = 'APCO'
    elif ie_type.find('Charging Id') != -1:
        ie_type = 'Charging ID'
    elif ie_type.find('H(e)NB Information Reporting') != -1:
        ie_type = 'eNB Information Reporting'
    elif ie_type.find('IPv4 Configuration Parameters (IP4CP)') != -1:
        ie_type = 'IP4CP'
    if ie_type not in type_list.keys():
        assert False, "Unknown IE type : [" \
                + cells[3].text + "]" + "(" + ie_type + ")"
    presence = cells[1].text.encode('ascii', 'ignore')
    ie_value = re.sub('\s*\n*\s*\([^\)]*\)*', '', cells[0].text).encode('ascii', 'ignore')
    comment = cells[2].text.encode('ascii', 'ignore')
    comment = re.sub('\n|\"|\'|\\\\', '', comment);

    if int(instance) > int(type_list[ie_type]["max_instance"]):
        type_list[ie_type]["max_instance"] = instance
        write_file(f, "type_list[\"" + ie_type + "\"][\"max_instance\"] = \"" + instance + "\"\n")

    return { "ie_type" : ie_type, "ie_value" : ie_value, "presence" : presence, "instance" : instance, "comment" : comment }

def write_cells_to_file(name, cells):
    write_file(f, name + ".append({ \"ie_type\" : \"" + cells["ie_type"] + \
        "\", \"ie_value\" : \"" + cells["ie_value"] + \
        "\", \"presence\" : \"" + cells["presence"] + \
        "\", \"instance\" : \"" + cells["instance"] + \
        "\", \"comment\" : \"" + cells["comment"] + "\"})\n")

try:
    opts, args = getopt.getopt(sys.argv[1:], "df:ho:c:", ["debug", "file", "help", "output", "cache"])
except getopt.GetoptError as err:
    # print help information and exit:
    usage()
    sys.exit(2)

for o, a in opts:
    if o in ("-d", "--debug"):
        verbosity = 1
    if o in ("-f", "--file"):
        filename = a
    if o in ("-o", "--output"):
        outdir = a
        if outdir.rfind('/') != len(outdir):
            outdir += '/'
    if o in ("-c", "--cache"):
        cache = a
        if cachedir.rfind('/') != len(cachedir):
            cachedir += '/'
    if o in ("-h", "--help"):
        usage()
        sys.exit(2)

if os.path.isfile(filename) and os.access(filename, os.R_OK):
    file = open(filename, 'r') 
else:
    d_error("Cannot find file : " + filename)

d_info("[Message List]")
cachefile = cachedir + 'tlv_msg_list.py'
if os.path.isfile(cachefile) and os.access(cachefile, os.R_OK):
    execfile(cachefile)
    print "Read from " + cachefile
else:
    document = Document(filename)
    f = open(cachefile, 'w') 

    msg_table = ""
    for i, table in enumerate(document.tables):
        cell = table.rows[0].cells[0]
        if cell.text.find('Message Type value') != -1:
            msg_table = table
            d_print("Table Index = %d\n" % i)

    for row in msg_table.rows[2:-4]:
        key = row.cells[1].text.encode('ascii', 'ignore')
        type = row.cells[0].text.encode('ascii', 'ignore')
        if type.isdigit() is False:
            continue
        if int(type) in range(128, 160):
            continue
        if int(type) in range(231, 240):
            continue
        if key.find('Reserved') != -1:
            continue
        key = re.sub('\s*\n*\s*\([^\)]*\)*', '', key)
        msg_list[key] = { "type": type }
        write_file(f, "msg_list[\"" + key + "\"] = { \"type\" : \"" + type + "\" }\n")
    f.close()

d_info("[IE Type List]")
cachefile = cachedir + 'tlv_type_list.py'
if os.path.isfile(cachefile) and os.access(cachefile, os.R_OK):
    execfile(cachefile)
    print "Read from " + cachefile
else:
    document = Document(filename)
    f = open(cachefile, 'w') 

    ie_table = ""
    for i, table in enumerate(document.tables):
        cell = table.rows[0].cells[0]
        if cell.text.find('IE Type value') != -1:
            ie_table = table
            d_print("Table Index = %d\n" % i)

    for row in ie_table.rows[1:-5]:
        key = row.cells[1].text.encode('ascii', 'ignore')
        if key.find('Reserved') != -1:
            continue
        if key.find('MM Context') != -1:
            continue
        elif key.find('Recovery') != -1:
            key = 'Recovery'
        elif key.find('Trusted WLAN Mode Indication') != -1:
            key = 'TWMI'
        elif key.find('LDN') != -1:
            key = 'LDN'
        elif key.find('APCO') != -1:
            key = 'APCO'
        elif key.find('Remote UE IP information') != -1:
            key = 'Remote UE IP Information'
        elif key.find('Procedure Transaction ID') != -1:
            key = 'PTI'
        else:
            key = re.sub('.*\(', '', row.cells[1].text.encode('ascii', 'ignore'))
            key = re.sub('\)', '', key)
            key = re.sub('\s*$', '', key)
        type = row.cells[0].text.encode('ascii', 'ignore')
        type_list[key] = { "type": type , "max_instance" : "0" }
        write_file(f, "type_list[\"" + key + "\"] = { \"type\" : \"" + type)
        write_file(f, "\", \"max_instance\" : \"0\" }\n")
    f.close()
type_list['MM Context'] = { "type": "107", "max_instance" : "0" }

d_info("[Group IE List]")
cachefile = cachedir + 'tlv_group_list.py'
if os.path.isfile(cachefile) and os.access(cachefile, os.R_OK):
    execfile(cachefile)
    print "Read from " + cachefile
else:
    document = Document(filename)
    f = open(cachefile, 'w') 

    for i, table in enumerate(document.tables):
        if table.rows[0].cells[0].text.find('Octet') != -1 and \
            table.rows[0].cells[2].text.find('IE Type') != -1:
            d_print("Table Index = %d\n" % i)

            row = table.rows[0];

            if len(re.findall('\d+', row.cells[2].text)) == 0:
                continue;
            ie_type = re.findall('\d+', row.cells[2].text)[0].encode('ascii', 'ignore')
            ie_name = re.sub('\s*IE Type.*', '', row.cells[2].text.encode('ascii', 'ignore'))

            if ie_name not in group_list.keys():
                ies = []
                write_file(f, "ies = []\n")
                for row in table.rows[4:]:
                    cells = get_cells(row.cells)
                    if cells is None:
                        continue

                    ies_is_added = True
                    for ie in ies:
                        if (cells["ie_type"], cells["instance"]) == (ie["ie_type"], ie["instance"]):
                            ies_is_added = False
                    if ies_is_added is True:
                        ies.append(cells)
                        write_cells_to_file("ies", cells)

                group_list[ie_name] = { "type" : ie_type, "ies" : ies }
                write_file(f, "group_list[\"" + ie_name + "\"] = { \"type\" : \"" + ie_type + "\", \"ies\" : ies }\n")
            else:
                group_list_is_added = False
                added_ies = group_list[ie_name]["ies"]
                write_file(f, "added_ies = group_list[\"" + ie_name + "\"][\"ies\"]\n")
                for row in table.rows[4:]:
                    cells = get_cells(row.cells)
                    if cells is None:
                        continue

                    ies_is_added = True
                    for ie in group_list[ie_name]["ies"]:
                        if (cells["ie_type"], cells["instance"]) == (ie["ie_type"], ie["instance"]):
                            ies_is_added = False
                    for ie in ies:
                        if (cells["ie_type"], cells["instance"]) == (ie["ie_type"], ie["instance"]):
                            ies_is_added = False
                    if ies_is_added is True:
                        added_ies.append(cells)
                        write_cells_to_file("added_ies", cells)
                        group_list_is_added = True
                if group_list_is_added is True:
                    group_list[ie_name] = { "type" : ie_type, "ies" : added_ies }
                    write_file(f, "group_list[\"" + ie_name + "\"] = { \"type\" : \"" + ie_type + "\", \"ies\" : added_ies }\n")
    f.close()

msg_list["Echo Request"]["table"] = 6
msg_list["Echo Response"]["table"] = 7
msg_list["Create Session Request"]["table"] = 8
msg_list["Create Session Response"]["table"] = 13
msg_list["Create Bearer Request"]["table"] = 18
msg_list["Create Bearer Response"]["table"] = 22
msg_list["Modify Bearer Request"]["table"] = 29
msg_list["Modify Bearer Response"]["table"] = 33
msg_list["Delete Session Request"]["table"] = 38
msg_list["Delete Bearer Request"]["table"] = 40
msg_list["Delete Session Response"]["table"] = 44
msg_list["Delete Bearer Response"]["table"] = 47
msg_list["Downlink Data Notification"]["table"] = 50
msg_list["Downlink Data Notification Acknowledge"]["table"] = 53
msg_list["Downlink Data Notification Failure Indication"]["table"] = 54
msg_list["Delete Indirect Data Forwarding Tunnel Request"]["table"] = 55
msg_list["Delete Indirect Data Forwarding Tunnel Response"]["table"] = 56
msg_list["Modify Bearer Command"]["table"] = 57
msg_list["Modify Bearer Failure Indication"]["table"] = 60
msg_list["Update Bearer Request"]["table"] = 62
msg_list["Update Bearer Response"]["table"] = 66
msg_list["Delete Bearer Command"]["table"] = 69
msg_list["Delete Bearer Failure Indication"]["table"] = 72
msg_list["Create Indirect Data Forwarding Tunnel Request"]["table"] = 75
msg_list["Create Indirect Data Forwarding Tunnel Response"]["table"] = 77
msg_list["Release Access Bearers Request"]["table"] = 79
msg_list["Release Access Bearers Response"]["table"] = 80
msg_list["Modify Access Bearers Request"]["table"] = 84
msg_list["Modify Access Bearers Response"]["table"] = 87

for key in msg_list.keys():
    if "table" in msg_list[key].keys():
        d_info("[" + key + "]")
        cachefile = cachedir + "tlv_msg_" + msg_list[key]["type"] + ".py"
        if os.path.isfile(cachefile) and os.access(cachefile, os.R_OK):
            execfile(cachefile)
            print "Read from " + cachefile
        else:
            document = Document(filename)
            f = open(cachefile, 'w') 

            ies = []
            write_file(f, "ies = []\n")
            table = document.tables[msg_list[key]["table"]]
            for row in table.rows[1:]:
                cells = get_cells(row.cells)
                if cells is None:
                    continue

                ies_is_added = True
                for ie in ies:
                    if (cells["ie_type"], cells["instance"]) == (ie["ie_type"], ie["instance"]):
                        ies_is_added = False
                if ies_is_added is True:
                    ies.append(cells)
                    write_cells_to_file("ies", cells)
            msg_list[key]["ies"] = ies
            write_file(f, "msg_list[key][\"ies\"] = ies\n")
            f.close()

type_list["Recovery"]["size"] = 1                       # Type : 3
type_list["EBI"]["size"] = 1                            # Type : 73
type_list["RAT Type"]["size"] = 1                       # Type : 82
type_list["PDN Type"]["size"] = 1                       # Type : 99
type_list["Port Number"]["size"] = 2                    # Type : 126
type_list["APN Restriction"]["size"] = 1                # Type : 127
type_list["Selection Mode"]["size"] = 1                 # Type : 128
type_list["Node Type"]["size"] = 1                 # Type : 128

f = open(outdir + 'gtp_message.h', 'w')
output_header_to_file(f)
f.write("""#ifndef __GTP_TLV_H__
#define __GTP_TLV_H__

#include "core_tlv_msg.h"

#ifdef __cplusplus
extern "C" {
#endif /* __cplusplus */

/* 5.1 General format */
#define GTPV1U_HEADER_LEN   8
#define GTPV2C_HEADER_LEN   12
#define GTP_TEID_LEN        4
typedef struct _gtp_header_t {
    union {
        struct {
        ED4(c_uint8_t version:3;,
            c_uint8_t piggybacked:1;,
            c_uint8_t teid_presence:1;,
            c_uint8_t spare1:3;)
        };
/* GTU-U flags */
#define GTPU_FLAGS_PN                       0x1
#define GTPU_FLAGS_S                        0x2
        c_uint8_t flags;
    };
    c_uint8_t type;
    c_uint16_t length;
    union {
        struct {
            c_uint32_t teid;
            /* sqn : 31bit ~ 8bit, spare : 7bit ~ 0bit */
#define GTP_XID_TO_SQN(__xid) htonl(((__xid) << 8))
#define GTP_SQN_TO_XID(__sqn) (ntohl(__sqn) >> 8)
            c_uint32_t sqn;
        };
        /* sqn : 31bit ~ 8bit, spare : 7bit ~ 0bit */
        c_uint32_t sqn_only;
    };
} __attribute__ ((packed)) gtp_header_t;

/* GTP-U message type, defined in 3GPP TS 29.281 Release 11 */
#define GTPU_MSGTYPE_ECHO_REQ               1
#define GTPU_MSGTYPE_ECHO_RSP               2
#define GTPU_MSGTYPE_ERR_IND                26
#define GTPU_MSGTYPE_SUPP_EXTHDR_NOTI       31
#define GTPU_MSGTYPE_END_MARKER             254
#define GTPU_MSGTYPE_GPDU                   255

/* GTPv2-C message type */
""")

tmp = [(k, v["type"]) for k, v in msg_list.items()]
sorted_msg_list = sorted(tmp, key=lambda tup: int(tup[1]))
for (k, v) in sorted_msg_list:
    f.write("#define GTP_" + v_upper(k) + "_TYPE " + v + "\n")
f.write("\n")

tmp = [(k, v["type"]) for k, v in type_list.items()]
sorted_type_list = sorted(tmp, key=lambda tup: int(tup[1]))
for (k, v) in sorted_type_list:
    f.write("#define TLV_" + v_upper(k) + "_TYPE " + v + "\n")
f.write("\n")

f.write("/* Infomration Element TLV Descriptor */\n")
for (k, v) in sorted_type_list:
    if k in group_list.keys():
        continue
    for instance in range(0, int(type_list[k]["max_instance"])+1):
        f.write("extern tlv_desc_t tlv_desc_" + v_lower(k))
        f.write("_" + str(instance) + ";\n")
f.write("\n")

tmp = [(k, v["type"]) for k, v in group_list.items()]
sorted_group_list = sorted(tmp, key=lambda tup: int(tup[1]))

f.write("/* Group Infomration Element TLV Descriptor */\n")
for (k, v) in sorted_group_list:
    for instance in range(0, int(type_list[k]["max_instance"])+1):
        f.write("extern tlv_desc_t tlv_desc_" + v_lower(k))
        f.write("_" + str(instance) + ";\n")
f.write("\n")

f.write("/* Message Descriptor */\n")
for (k, v) in sorted_msg_list:
    f.write("extern tlv_desc_t tlv_desc_" + v_lower(k) + ";\n")
f.write("\n")

f.write("/* Structure for Infomration Element */\n")
for (k, v) in sorted_type_list:
    if k in group_list.keys():
        continue
    if "size" in type_list[k]:
        if type_list[k]["size"] == 1:
            f.write("typedef tlv_uint8_t tlv_" + v_lower(k) + "_t;\n")
        elif type_list[k]["size"] == 2:
            f.write("typedef tlv_uint16_t tlv_" + v_lower(k) + "_t;\n")
        elif type_list[k]["size"] == 3:
            f.write("typedef tlv_uint24_t tlv_" + v_lower(k) + "_t;\n")
        elif type_list[k]["size"] == 4:
            f.write("typedef tlv_uint32_t tlv_" + v_lower(k) + "_t;\n")
        else:
            assert False, "Unknown size = %d for key = %s" % (type_list[k]["size"], k)
    else:
        f.write("typedef tlv_octet_t tlv_" + v_lower(k) + "_t;\n")
f.write("\n")

f.write("/* Structure for Group Infomration Element */\n")
for (k, v) in sorted_group_list:
    f.write("typedef struct _tlv_" + v_lower(k) + "_t {\n")
    f.write("    tlv_presence_t presence;\n")
    for ies in group_list[k]["ies"]:
        f.write("    tlv_" + v_lower(ies["ie_type"]) + "_t " + \
                v_lower(ies["ie_value"]))
        if ies["ie_type"] == "F-TEID":
            if ies["ie_value"] == "S2b-U ePDG F-TEID":
                f.write("_" + ies["instance"] + ";")
            elif ies["ie_value"] == "S2a-U TWAN F-TEID":
                f.write("_" + ies["instance"] + ";")
            else:
                f.write(";")
            f.write(" /* Instance : " + ies["instance"] + " */\n")
        else:
            f.write(";\n")
    f.write("} tlv_" + v_lower(k) + "_t;\n")
    f.write("\n")

f.write("/* Structure for Message */\n")
for (k, v) in sorted_msg_list:
    if "ies" in msg_list[k]:
        f.write("typedef struct _gtp_" + v_lower(k) + "_t {\n")
        for ies in msg_list[k]["ies"]:
            f.write("    tlv_" + v_lower(ies["ie_type"]) + "_t " + \
                    v_lower(ies["ie_value"]) + ";\n")
        f.write("} gtp_" + v_lower(k) + "_t;\n")
        f.write("\n")

f.write("typedef struct _gtp_message_t {\n")
f.write("   gtp_header_t h;\n")
f.write("   union {\n")
for (k, v) in sorted_msg_list:
    if "ies" in msg_list[k]:
        f.write("        gtp_" + v_lower(k) + "_t " + v_lower(k) + ";\n");
f.write("   };\n");
f.write("} gtp_message_t;\n\n")

f.write("""CORE_DECLARE(status_t) gtp_parse_msg(
        gtp_message_t *gtp_message, pkbuf_t *pkbuf);
CORE_DECLARE(status_t) gtp_build_msg(
        pkbuf_t **pkbuf, gtp_message_t *gtp_message);

#ifdef __cplusplus
}
#endif /* __cplusplus */

#endif /* __GTP_TLV_H__ */
""")
f.close()

f = open(outdir + 'gtp_message.c', 'w')
output_header_to_file(f)
f.write("""#define TRACE_MODULE _gtp_message
#include "core_debug.h"
#include "gtp_message.h"

""")

for (k, v) in sorted_type_list:
    if k in group_list.keys():
        continue
    for instance in range(0, int(type_list[k]["max_instance"])+1):
        f.write("tlv_desc_t tlv_desc_%s_%d =\n" % (v_lower(k), instance))
        f.write("{\n")
        if "size" in type_list[k]:
            if type_list[k]["size"] == 1:
                f.write("    TLV_UINT8,\n")
            elif type_list[k]["size"] == 2:
                f.write("    TLV_UINT16,\n")
            elif type_list[k]["size"] == 3:
                f.write("    TLV_UINT24,\n")
            elif type_list[k]["size"] == 4:
                f.write("    TLV_UINT32,\n")
            else:
                assert False, "Unknown size = %d for key = %s" % (type_list[k]["size"], k)
        else:
            f.write("    TLV_VAR_STR,\n")
        f.write("    \"%s\",\n" % k)
        f.write("    TLV_%s_TYPE,\n" % v_upper(k))
        if "size" in type_list[k]:
            f.write("    %d,\n" % type_list[k]["size"])
        else:
            f.write("    0,\n")
        f.write("    %d,\n" % instance)
        f.write("    sizeof(tlv_%s_t),\n" % v_lower(k))
        f.write("    { NULL }\n")
        f.write("};\n\n")

for (k, v) in sorted_group_list:
    for instance in range(0, int(type_list[k]["max_instance"])+1):
        f.write("tlv_desc_t tlv_desc_%s_%d =\n" % (v_lower(k), instance))
        f.write("{\n")
        f.write("    TLV_COMPOUND,\n")
        f.write("    \"%s\",\n" % k)
        f.write("    TLV_%s_TYPE,\n" % v_upper(k))
        f.write("    0,\n")
        f.write("    %d,\n" % instance)
        f.write("    sizeof(tlv_%s_t),\n" % v_lower(k))
        f.write("    {\n")
        for ies in group_list[k]["ies"]:
                f.write("        &tlv_desc_%s_%s,\n" % (v_lower(ies["ie_type"]), v_lower(ies["instance"])))
        f.write("        NULL,\n")
        f.write("    }\n")
        f.write("};\n\n")

for (k, v) in sorted_msg_list:
    if "ies" in msg_list[k]:
        f.write("tlv_desc_t tlv_desc_%s =\n" % v_lower(k))
        f.write("{\n")
        f.write("    TLV_MESSAGE,\n")
        f.write("    \"%s\",\n" % k)
        f.write("    0, 0, 0, 0, {\n")
        for ies in msg_list[k]["ies"]:
                f.write("        &tlv_desc_%s_%s,\n" % (v_lower(ies["ie_type"]), v_lower(ies["instance"])))
        f.write("    NULL,\n")
        f.write("}};\n\n")
f.write("\n")

f.write("""status_t gtp_parse_msg(gtp_message_t *gtp_message, pkbuf_t *pkbuf)
{
    status_t rv = CORE_ERROR;
    gtp_header_t *h = NULL;
    c_uint16_t size = 0;

    d_assert(gtp_message, return CORE_ERROR, "Null param");
    d_assert(pkbuf, return CORE_ERROR, "Null param");
    d_assert(pkbuf->payload, return CORE_ERROR, "Null param");

    d_trace(50, "[GTPv2] RECV : ");
    d_trace_hex(50, pkbuf->payload, pkbuf->len);

    h = pkbuf->payload;
    d_assert(h, return CORE_ERROR, "Null param");
    
    memset(gtp_message, 0, sizeof(gtp_message_t));

    if (h->teid_presence)
        size = GTPV2C_HEADER_LEN;
    else
        size = GTPV2C_HEADER_LEN-GTP_TEID_LEN;

    d_assert(pkbuf_header(pkbuf, -size) == CORE_OK,
            return CORE_ERROR, "pkbuf_header error");
    memcpy(&gtp_message->h, pkbuf->payload - size, size);

    if (h->teid_presence)
        gtp_message->h.teid = ntohl(gtp_message->h.teid);

    if (pkbuf->len == 0)
        return CORE_OK;

    switch(gtp_message->h.type)
    {
""")
for (k, v) in sorted_msg_list:
    if "ies" in msg_list[k]:
        f.write("        case GTP_%s_TYPE:\n" % v_upper(k))
        f.write("            rv = tlv_parse_msg(&gtp_message->%s,\n" % v_lower(k))
        f.write("                    &tlv_desc_%s, pkbuf, TLV_MODE_T1_L2_I1);\n" % v_lower(k))
        f.write("            break;\n")
f.write("""        default:
            d_warn("Not implmeneted(type:%d)", gtp_message->h.type);
            break;
    }

    return rv;
}

""")

f.write("""status_t gtp_build_msg(pkbuf_t **pkbuf, gtp_message_t *gtp_message)
{
    status_t rv = CORE_ERROR;

    d_assert(gtp_message, return rv, "Null param");
    switch(gtp_message->h.type)
    {
""")
for (k, v) in sorted_msg_list:
    if "ies" in msg_list[k]:
        f.write("        case GTP_%s_TYPE:\n" % v_upper(k))
        f.write("            rv = tlv_build_msg(pkbuf, &tlv_desc_%s,\n" % v_lower(k))
        f.write("                    &gtp_message->%s, TLV_MODE_T1_L2_I1);\n" % v_lower(k))
        f.write("            break;\n")
f.write("""        default:
            d_warn("Not implmeneted(type:%d)", gtp_message->h.type);
            break;
    }

    if ((*pkbuf) && (*pkbuf)->payload)
    {
        d_trace(50, "[GTPv2] SEND : ");
        d_trace_hex(50, (*pkbuf)->payload, (*pkbuf)->len);
    }

    return rv;
}
""")

f.write("\n")

f.close()
